import os
from decimal import Decimal

from django.conf import settings
from django.core.files import File
from docx import Document

from documents.models import InvoiceTemplate
import logging


logger = logging.getLogger(__name__)


def replace_text_in_paragraphs(document, replacements):
    for paragraph in document.paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                paragraph.text = paragraph.text.replace(old_text, str(new_text))


def find_invoice_table(document):
    for table in document.tables:
        first_row_text = [cell.text.strip() for cell in table.rows[0].cells]

        if "№" in first_row_text and "Наименование" in first_row_text:
            return table
    logger.error('Не найдена таблица в шаблоне счета')
    raise ValueError("Не найдена таблица товаров в шаблоне счёта.")


def fill_invoice_table(document, invoice):
    table = find_invoice_table(document)

    # удаляем все строки кроме заголовка
    while len(table.rows) > 1:
        row = table.rows[1]
        table._tbl.remove(row._tr)

    total = Decimal("0.00")

    for index, item in enumerate(invoice.items.all(), start=1):
        row_cells = table.add_row().cells

        item_total = item.total
        total += item_total

        row_cells[0].text = str(index)
        row_cells[1].text = item.product_name
        row_cells[2].text = item.get_unit_display()
        row_cells[3].text = str(item.quantity)
        row_cells[4].text = str(item.unit_price)
        row_cells[5].text = str(item_total)

    return total


def generate_invoice_docx(invoice):
    logger.info(f'Начало генерации файла Invoice id={invoice.id}')
    template = (
        InvoiceTemplate.objects.filter(is_active=True).order_by("-created_at").first()
    )

    if not template:
        logger.error('Нет активного шаблона InvoiceTemplate')
        raise ValueError("Не найден активный шаблон счёта.")

    document = Document(template.file.path)

    total = fill_invoice_table(document, invoice)
    vat = total * Decimal("0.22")
    total_with_vat = total + vat

    replacements = {
        "{{ number }}": invoice.number or invoice.pk,
        "{{ date }}": invoice.created_at.strftime("%d.%m.%Y"),
        "{{ client_name }}": invoice.client.name,
        "{{ total }}": f"{total:.2f}",
        "{{ vat }}": f"{vat:.2f}",
        "{{ total_with_vat }}": f"{total_with_vat:.2f}",
    }

    replace_text_in_paragraphs(document, replacements)

    file_name = f"invoice_{invoice.pk}.docx"
    file_dir = os.path.join(settings.MEDIA_ROOT, "invoices")
    file_path = os.path.join(file_dir, file_name)

    os.makedirs(file_dir, exist_ok=True)

    document.save(file_path)

    with open(file_path, "rb") as docx_file:
        invoice.file.save(file_name, File(docx_file), save=True)
        logger.info(f'Файл Invoice id={invoice.id} создан')

    return invoice.file
